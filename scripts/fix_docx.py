#!/usr/bin/env python3
"""Post-process DOCX: add formatted title page and automatic table of contents (ЗМІСТ)."""

import sys
import os
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_run_font(run, font_name='Times New Roman'):
    """Set font on all character ranges for a run."""
    rPr = run._r.find(qn('w:rPr'))
    if rPr is None:
        rPr = OxmlElement('w:rPr')
        run._r.insert(0, rPr)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    for attr in ['ascii', 'hAnsi', 'cs', 'eastAsia']:
        rFonts.set(qn(f'w:{attr}'), font_name)
    for attr in ['asciiTheme', 'hAnsiTheme', 'cstheme', 'eastAsiaTheme']:
        key = qn(f'w:{attr}')
        if key in rFonts.attrib:
            del rFonts.attrib[key]


def set_style_font(style, font_name='Times New Roman'):
    """Set font on all character ranges for a style."""
    rPr = style.element.find(qn('w:rPr'))
    if rPr is None:
        rPr = OxmlElement('w:rPr')
        style.element.append(rPr)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    for attr in ['ascii', 'hAnsi', 'cs', 'eastAsia']:
        rFonts.set(qn(f'w:{attr}'), font_name)
    for attr in ['asciiTheme', 'hAnsiTheme', 'cstheme', 'eastAsiaTheme']:
        key = qn(f'w:{attr}')
        if key in rFonts.attrib:
            del rFonts.attrib[key]


def configure_toc_styles(doc):
    """Create/update TOC 1/2/3 styles with right tab stop + dot leader.

    Text width = 210mm - 30mm(left) - 10mm(right) = 170mm ≈ 9638 twips.
    """
    text_width_twips = 9638

    for level, indent_cm in [('TOC 1', 0), ('TOC 2', 1.0), ('TOC 3', 2.0)]:
        if level not in [s.name for s in doc.styles]:
            toc_style = doc.styles.add_style(level, 1)  # paragraph style
        else:
            toc_style = doc.styles[level]

        toc_style.font.name = 'Times New Roman'
        toc_style.font.size = Pt(14)
        set_style_font(toc_style, 'Times New Roman')
        toc_style.paragraph_format.line_spacing = 1.5
        toc_style.paragraph_format.space_before = Pt(0)
        toc_style.paragraph_format.space_after = Pt(0)
        toc_style.paragraph_format.first_line_indent = Cm(0)
        if indent_cm > 0:
            toc_style.paragraph_format.left_indent = Cm(indent_cm)

        # Ensure pPr exists
        pPr = toc_style.element.find(qn('w:pPr'))
        if pPr is None:
            pPr = OxmlElement('w:pPr')
            toc_style.element.append(pPr)

        # Remove any existing tabs
        existing_tabs = pPr.find(qn('w:tabs'))
        if existing_tabs is not None:
            pPr.remove(existing_tabs)

        # Add right tab stop with dot leader
        tabs = OxmlElement('w:tabs')
        tab = OxmlElement('w:tab')
        tab.set(qn('w:val'), 'right')
        tab.set(qn('w:leader'), 'dot')
        tab.set(qn('w:pos'), str(text_width_twips))
        tabs.append(tab)
        pPr.append(tabs)


def add_para(doc, text='', alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=False,
             font_size=Pt(14), space_before=Pt(0), space_after=Pt(0)):
    """Add a formatted paragraph to the document."""
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = space_before
    p.paragraph_format.space_after = space_after
    p.paragraph_format.line_spacing = 1.2
    if text:
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = font_size
        run.font.bold = bold
        set_run_font(run, 'Times New Roman')
    return p


def add_page_break(doc):
    """Add a page break paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)
    return p


def add_toc_field(doc):
    """Add a Word TOC field that auto-generates table of contents."""
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)

    # Begin field
    run1 = p.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    run1._r.append(fldChar1)

    # Field instruction
    run2 = p.add_run()
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = ' TOC \\o "1-3" \\h \\z \\u '
    run2._r.append(instrText)

    # Separator
    run3 = p.add_run()
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    run3._r.append(fldChar2)

    # Placeholder (replaced automatically when Word opens the document)
    run4 = p.add_run('(Зміст оновиться автоматично)')
    run4.font.name = 'Times New Roman'
    run4.font.size = Pt(14)
    set_run_font(run4, 'Times New Roman')

    # End field
    run5 = p.add_run()
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    run5._r.append(fldChar3)

    return p


def build_title_page(doc, src_dir):
    """Build title page elements. Returns list of XML elements to move to beginning."""
    elements = []

    # --- Header lines ---
    p = add_para(doc, 'МІНІСТЕРСТВО ОСВІТИ І НАУКИ УКРАЇНИ', space_after=Pt(6))
    elements.append(p._element)

    p = add_para(doc, 'НАЦІОНАЛЬНИЙ УНІВЕРСИТЕТ «ЛЬВІВСЬКА ПОЛІТЕХНІКА»', space_after=Pt(6))
    elements.append(p._element)

    p = add_para(doc, 'КАФЕДРА ЕЛЕКТРОННОЇ ІНЖЕНЕРІЇ', space_after=Pt(24))
    elements.append(p._element)

    # --- University logo ---
    logo_path = os.path.join(src_dir, 'univ-logo.png')
    if os.path.exists(logo_path):
        doc.add_picture(logo_path, width=Cm(6))
        pic_para = doc.paragraphs[-1]
        pic_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pic_para.paragraph_format.first_line_indent = Cm(0)
        pic_para.paragraph_format.space_before = Pt(12)
        pic_para.paragraph_format.space_after = Pt(24)
        elements.append(pic_para._element)

    # --- Title block ---
    p = add_para(doc, 'ДИПЛОМНА РОБОТА', bold=True, font_size=Pt(16), space_after=Pt(8))
    elements.append(p._element)

    p = add_para(doc, 'на здобуття ступеня вищої освіти', space_after=Pt(6))
    elements.append(p._element)

    p = add_para(doc, 'БАКАЛАВР', bold=True, space_after=Pt(12))
    elements.append(p._element)

    p = add_para(doc, 'за спеціальністю G5 «Електроніка»', bold=True, space_after=Pt(36))
    elements.append(p._element)

    # --- Author / supervisor block (right-aligned) ---
    author_lines = [
        'Виконав:',
        'студент 4 курсу',
        'групи ЕЛ-41',
        'Воронов Сергій Олексійович',
        '',
        'Керівник:',
        'доктор філософії, доцент',
        'Мельников Сергій Олександрович',
    ]
    for line in author_lines:
        p = add_para(doc, line, alignment=WD_ALIGN_PARAGRAPH.RIGHT, space_after=Pt(2))
        elements.append(p._element)

    # --- City and year (pushed toward bottom with large space_before) ---
    p = add_para(doc, 'Львів \u2014 2026', space_before=Pt(96))
    elements.append(p._element)

    # --- Page break after title page ---
    p = add_page_break(doc)
    elements.append(p._element)

    return elements


def build_toc_page(doc):
    """Build table of contents page. Returns list of XML elements."""
    elements = []

    # ЗМІСТ heading
    p = add_para(doc, 'ЗМІСТ', bold=True, space_after=Pt(12))
    elements.append(p._element)

    # TOC field
    p = add_toc_field(doc)
    elements.append(p._element)

    # Page break after TOC
    p = add_page_break(doc)
    elements.append(p._element)

    return elements


def fix_docx(input_path, output_path, src_dir):
    """Add formatted title page and TOC to DOCX."""
    doc = Document(input_path)
    body = doc.element.body

    # Configure TOC styles with dot leaders (must be in output doc, not just reference)
    configure_toc_styles(doc)

    # Tell Word to auto-update all fields (including TOC) when document opens
    settings = doc.settings.element
    update_fields = settings.find(qn('w:updateFields'))
    if update_fields is None:
        update_fields = OxmlElement('w:updateFields')
        settings.append(update_fields)
    update_fields.set(qn('w:val'), 'true')

    # Build title page and TOC elements (added at end of document)
    title_elements = build_title_page(doc, src_dir)
    toc_elements = build_toc_page(doc)

    all_new = title_elements + toc_elements

    # Move all new elements to the beginning of the document
    for i, elem in enumerate(all_new):
        body.remove(elem)
        body.insert(i, elem)

    doc.save(output_path)
    print(f"  Title page and TOC added: {output_path}")


if __name__ == '__main__':
    if len(sys.argv) < 3:
        print(f"Usage: {sys.argv[0]} <input.docx> <output.docx> [--src-dir=<src_dir>]")
        sys.exit(1)

    input_path = sys.argv[1]
    output_path = sys.argv[2]

    src_dir = 'src'
    for arg in sys.argv[3:]:
        if arg.startswith('--src-dir='):
            src_dir = arg.split('=', 1)[1]

    fix_docx(input_path, output_path, src_dir)
