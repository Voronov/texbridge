#!/usr/bin/env python3
"""Create a reference DOCX template with Ukrainian university standard formatting (ДСТУ 3008:2015)."""

from docx import Document
from docx.shared import Pt, Cm, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
import sys


def set_font_all_ranges(style, font_name):
    """Set font for ALL character ranges and remove Word theme font overrides.

    python-docx's font.name only sets the ASCII range. Word defaults to
    Calibri for other ranges via theme fonts. This sets all four ranges
    explicitly and removes theme references.
    """
    rPr = style.element.find(qn('w:rPr'))
    if rPr is None:
        rPr = style.element.makeelement(qn('w:rPr'), {})
        style.element.append(rPr)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = rPr.makeelement(qn('w:rFonts'), {})
        rPr.append(rFonts)
    # Set explicit font for all ranges
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rFonts.set(qn('w:cs'), font_name)
    rFonts.set(qn('w:eastAsia'), font_name)
    # Remove theme font references that override explicit fonts
    for attr in ['asciiTheme', 'hAnsiTheme', 'cstheme', 'eastAsiaTheme']:
        key = qn(f'w:{attr}')
        if key in rFonts.attrib:
            del rFonts.attrib[key]


def remove_heading_color(style):
    """Remove color overrides from heading styles (Word defaults to blue)."""
    rpr = style.element.find(qn('w:rPr'))
    if rpr is not None:
        color = rpr.find(qn('w:color'))
        if color is not None:
            rpr.remove(color)


def create_reference_doc(output_path):
    doc = Document()

    # --- Page setup ---
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.left_margin = Mm(30)
    section.right_margin = Mm(10)
    section.top_margin = Mm(20)
    section.bottom_margin = Mm(20)
    section.header_distance = Mm(10)
    section.footer_distance = Mm(10)

    # --- Page numbering: top right ---
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = hp.add_run()
    # Add page number field
    fldChar1 = run._r.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'begin'})
    run._r.append(fldChar1)
    run2 = hp.add_run()
    instrText = run2._r.makeelement(qn('w:instrText'), {})
    instrText.text = ' PAGE '
    run2._r.append(instrText)
    run3 = hp.add_run()
    fldChar2 = run3._r.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'end'})
    run3._r.append(fldChar2)
    # Set header font
    for run in hp.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)

    # --- Default paragraph style (Body Text / Normal) ---
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(14)
    set_font_all_ranges(style, 'Times New Roman')
    pf = style.paragraph_format
    pf.line_spacing = 1.5
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.first_line_indent = Cm(1.25)
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # --- Body Text style (pandoc uses this) ---
    if 'Body Text' not in [s.name for s in doc.styles]:
        body_style = doc.styles.add_style('Body Text', 1)  # 1 = paragraph
    else:
        body_style = doc.styles['Body Text']
    body_style.font.name = 'Times New Roman'
    body_style.font.size = Pt(14)
    set_font_all_ranges(body_style, 'Times New Roman')
    body_style.paragraph_format.line_spacing = 1.5
    body_style.paragraph_format.space_before = Pt(0)
    body_style.paragraph_format.space_after = Pt(0)
    body_style.paragraph_format.first_line_indent = Cm(1.25)
    body_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # --- First Paragraph style ---
    if 'First Paragraph' not in [s.name for s in doc.styles]:
        fp_style = doc.styles.add_style('First Paragraph', 1)
    else:
        fp_style = doc.styles['First Paragraph']
    fp_style.font.name = 'Times New Roman'
    fp_style.font.size = Pt(14)
    set_font_all_ranges(fp_style, 'Times New Roman')
    fp_style.paragraph_format.line_spacing = 1.5
    fp_style.paragraph_format.space_before = Pt(0)
    fp_style.paragraph_format.space_after = Pt(0)
    fp_style.paragraph_format.first_line_indent = Cm(1.25)
    fp_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # --- Heading 1: РОЗДІЛ (bold, ALL CAPS, centered, 14pt) ---
    h1 = doc.styles['Heading 1']
    h1.font.name = 'Times New Roman'
    h1.font.size = Pt(14)
    h1.font.bold = True
    h1.font.color.rgb = None
    set_font_all_ranges(h1, 'Times New Roman')
    remove_heading_color(h1)
    h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h1.paragraph_format.space_before = Pt(0)
    h1.paragraph_format.space_after = Pt(12)
    h1.paragraph_format.first_line_indent = Cm(0)
    h1.paragraph_format.line_spacing = 1.5
    h1.paragraph_format.page_break_before = True

    # --- Heading 2: Підрозділ (bold, left with indent, 14pt) ---
    h2 = doc.styles['Heading 2']
    h2.font.name = 'Times New Roman'
    h2.font.size = Pt(14)
    h2.font.bold = True
    h2.font.color.rgb = None
    set_font_all_ranges(h2, 'Times New Roman')
    remove_heading_color(h2)
    h2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after = Pt(6)
    h2.paragraph_format.first_line_indent = Cm(1.25)
    h2.paragraph_format.line_spacing = 1.5

    # --- Heading 3: Пункт (bold, left with indent, 14pt) ---
    h3 = doc.styles['Heading 3']
    h3.font.name = 'Times New Roman'
    h3.font.size = Pt(14)
    h3.font.bold = True
    h3.font.color.rgb = None
    set_font_all_ranges(h3, 'Times New Roman')
    remove_heading_color(h3)
    h3.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h3.paragraph_format.space_before = Pt(6)
    h3.paragraph_format.space_after = Pt(6)
    h3.paragraph_format.first_line_indent = Cm(1.25)
    h3.paragraph_format.line_spacing = 1.5

    # --- Heading 4 (if exists) ---
    if 'Heading 4' in [s.name for s in doc.styles]:
        h4 = doc.styles['Heading 4']
        h4.font.name = 'Times New Roman'
        h4.font.size = Pt(14)
        h4.font.bold = False
        h4.font.color.rgb = None
        set_font_all_ranges(h4, 'Times New Roman')
        remove_heading_color(h4)
        h4.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        h4.paragraph_format.first_line_indent = Cm(1.25)
        h4.paragraph_format.line_spacing = 1.5

    # --- Caption style (for figures/tables) ---
    if 'Caption' not in [s.name for s in doc.styles]:
        cap_style = doc.styles.add_style('Caption', 1)
    else:
        cap_style = doc.styles['Caption']
    cap_style.font.name = 'Times New Roman'
    cap_style.font.size = Pt(14)
    cap_style.font.italic = False
    cap_style.font.color.rgb = None
    set_font_all_ranges(cap_style, 'Times New Roman')
    cap_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_style.paragraph_format.space_before = Pt(6)
    cap_style.paragraph_format.space_after = Pt(6)
    cap_style.paragraph_format.first_line_indent = Cm(0)
    cap_style.paragraph_format.line_spacing = 1.5

    # --- Figure style (centered images) ---
    if 'Figure' not in [s.name for s in doc.styles]:
        fig_style = doc.styles.add_style('Figure', 1)
    else:
        fig_style = doc.styles['Figure']
    fig_style.font.name = 'Times New Roman'
    fig_style.font.size = Pt(14)
    set_font_all_ranges(fig_style, 'Times New Roman')
    fig_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fig_style.paragraph_format.first_line_indent = Cm(0)

    # --- Table of Contents heading ---
    if 'TOC Heading' in [s.name for s in doc.styles]:
        toc = doc.styles['TOC Heading']
        toc.font.name = 'Times New Roman'
        toc.font.size = Pt(14)
        toc.font.bold = True
        set_font_all_ranges(toc, 'Times New Roman')
        toc.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # --- Source Code style (for verbatim/code blocks) ---
    if 'Source Code' not in [s.name for s in doc.styles]:
        sc_style = doc.styles.add_style('Source Code', 1)  # paragraph style
    else:
        sc_style = doc.styles['Source Code']
    sc_style.font.name = 'Courier New'
    sc_style.font.size = Pt(10)
    set_font_all_ranges(sc_style, 'Courier New')
    sc_style.paragraph_format.line_spacing = 1.0
    sc_style.paragraph_format.space_before = Pt(0)
    sc_style.paragraph_format.space_after = Pt(0)
    sc_style.paragraph_format.first_line_indent = Cm(0)
    sc_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    # Add border/shading via XML for visual distinction
    pPr = sc_style.element.find(qn('w:pPr'))
    if pPr is None:
        pPr = sc_style.element.makeelement(qn('w:pPr'), {})
        sc_style.element.append(pPr)
    # Light grey background shading
    shd = pPr.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear',
        qn('w:color'): 'auto',
        qn('w:fill'): 'F2F2F2'
    })
    pPr.append(shd)
    # Left indent to match body text
    ind = pPr.makeelement(qn('w:ind'), {
        qn('w:left'): '709',  # 1.25cm in twips
    })
    pPr.append(ind)

    # --- Verbatim Char style (inline code character style) ---
    if 'Verbatim Char' not in [s.name for s in doc.styles]:
        vc_style = doc.styles.add_style('Verbatim Char', 2)  # 2 = character style
    else:
        vc_style = doc.styles['Verbatim Char']
    vc_style.font.name = 'Courier New'
    vc_style.font.size = Pt(10)

    # --- TOC entry styles (with dot leaders) ---
    # Text width: 210mm - 30mm(left) - 10mm(right) = 170mm
    # In twips: 170mm * (1440/25.4) ≈ 9638
    text_width_twips = 9638
    for level, indent_cm in [('TOC 1', 0), ('TOC 2', 1.0), ('TOC 3', 2.0)]:
        if level not in [s.name for s in doc.styles]:
            toc_entry = doc.styles.add_style(level, 1)
        else:
            toc_entry = doc.styles[level]
        toc_entry.font.name = 'Times New Roman'
        toc_entry.font.size = Pt(14)
        set_font_all_ranges(toc_entry, 'Times New Roman')
        toc_entry.paragraph_format.line_spacing = 1.5
        toc_entry.paragraph_format.space_before = Pt(0)
        toc_entry.paragraph_format.space_after = Pt(0)
        toc_entry.paragraph_format.first_line_indent = Cm(0)
        if indent_cm > 0:
            toc_entry.paragraph_format.left_indent = Cm(indent_cm)
        # Right tab stop with dot leader for page numbers
        pPr_toc = toc_entry.element.find(qn('w:pPr'))
        if pPr_toc is None:
            pPr_toc = toc_entry.element.makeelement(qn('w:pPr'), {})
            toc_entry.element.append(pPr_toc)
        tabs = pPr_toc.makeelement(qn('w:tabs'), {})
        tab = tabs.makeelement(qn('w:tab'), {
            qn('w:val'): 'right',
            qn('w:leader'): 'dot',
            qn('w:pos'): str(text_width_twips),
        })
        tabs.append(tab)
        pPr_toc.append(tabs)

    # --- List styles ---
    for list_style_name in ['List Paragraph', 'List Bullet', 'List Number']:
        if list_style_name in [s.name for s in doc.styles]:
            ls = doc.styles[list_style_name]
            ls.font.name = 'Times New Roman'
            ls.font.size = Pt(14)
            set_font_all_ranges(ls, 'Times New Roman')
            ls.paragraph_format.line_spacing = 1.5
            ls.paragraph_format.space_before = Pt(0)
            ls.paragraph_format.space_after = Pt(0)

    # Add a sample paragraph to ensure styles are applied
    p = doc.add_paragraph('', style='Normal')
    p.clear()

    doc.save(output_path)
    print(f"  Reference DOCX created: {output_path}")


if __name__ == '__main__':
    output = sys.argv[1] if len(sys.argv) > 1 else 'reference.docx'
    create_reference_doc(output)
